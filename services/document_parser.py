"""
Multi-format document parser for PDF, DOCX, TXT, and Excel files.
Supports offline operation with enhanced metadata and citation tracking.
Includes multimodal support with image extraction and vision-based descriptions.
"""
import os
import logging
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from datetime import datetime
from langchain.schema import Document
import io

logger = logging.getLogger(__name__)

# PDF parsing
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
    PDF_LIBRARY = "pymupdf"
except ImportError:
    try:
        import pdfplumber
        PDF_AVAILABLE = True
        PDF_LIBRARY = "pdfplumber"
    except ImportError:
        PDF_AVAILABLE = False
        PDF_LIBRARY = None
        logger.warning("PDF parsing not available. Install PyMuPDF (pip install pymupdf) or pdfplumber.")

# DOCX parsing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("DOCX parsing not available. Install python-docx (pip install python-docx).")

# Excel parsing (already available via pandas)
import pandas as pd

# Image processing
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    logger.warning("PIL/Pillow not available. Image extraction will be limited.")


def parse_pdf(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200, use_ocr: bool = False) -> List[Document]:
    """
    Parse PDF file with page-level tracking and smart chunking.
    Supports both native PDFs and scanned PDFs (via OCR).
    
    Args:
        file_path: Path to PDF file
        chunk_size: Number of characters per chunk
        chunk_overlap: Overlap between chunks for context preservation
        use_ocr: If True, use OCR to extract text from scanned PDF
    
    Returns:
        List of Document objects with metadata including page numbers for citations
    """
    if not PDF_AVAILABLE:
        raise ImportError("PDF parsing libraries not installed. Install: pip install pymupdf")
    
    documents = []
    file_name = os.path.basename(file_path)
    
    # If OCR mode is enabled, use Tesseract OCR service
    if use_ocr:
        from services.ocr_service import ocr_pdf_document, is_tesseract_installed, extract_text_with_ocr
        
        if not is_tesseract_installed():
            raise ImportError("Tesseract OCR not installed. Install: apt-get install tesseract-ocr (Linux) or brew install tesseract (Mac)")
        
        logger.info(f"🔍 Using Tesseract OCR mode for: {file_name}")
        
        # Open PDF to extract both page text (via OCR) and images
        doc = fitz.open(file_path)
        total_pages = len(doc)
        
        # Process PDF with Tesseract OCR
        ocr_result = ocr_pdf_document(file_path)
        
        if not ocr_result.get("success"):
            error_msg = ocr_result.get("error", "Unknown OCR error")
            logger.error(f"Tesseract OCR failed: {error_msg}")
            doc.close()
            raise RuntimeError(f"Tesseract OCR processing failed: {error_msg}")
        
        logger.info(f"✓ Tesseract OCR completed: {ocr_result['processed_pages']} pages, avg confidence: {ocr_result['avg_confidence']}%, engine: {ocr_result.get('ocr_engine', 'Tesseract')}")
        
        # Process OCR results into documents (text from pages)
        total_ocr_text = 0
        for page_result in ocr_result["pages"]:
            page_num = page_result["page_number"]
            success = page_result.get("success", False)
            text = page_result.get("text", "").strip() if success else ""
            char_count = page_result.get("char_count", 0)
            
            logger.info(f"Page {page_num}: OCR success={success}, chars={char_count}, text_preview={text[:50] if text else 'N/A'}...")
            
            if success and text:
                total_ocr_text += len(text)
                confidence = page_result.get("confidence", 0)
                
                # Chunk the OCR text
                chunks = _chunk_text(text, chunk_size, chunk_overlap)
                logger.debug(f"Page {page_num}: Created {len(chunks)} chunks from {len(text)} chars")
                
                for chunk_idx, chunk in enumerate(chunks):
                    if chunk and chunk.strip():  # Only add non-empty chunks
                        documents.append(Document(
                            page_content=chunk,
                            metadata={
                                "source": file_name,
                                "file_path": file_path,
                                "page_number": page_num,
                                "chunk_index": chunk_idx,
                                "total_chunks_in_page": len(chunks),
                                "file_type": "pdf",
                                "ocr_applied": True,
                                "ocr_engine": "Tesseract",
                                "ocr_confidence": confidence,
                                "ingestion_timestamp": datetime.now().isoformat(),
                                "total_pages": total_pages,
                                "citation": f"{file_name}, Page {page_num} (OCR)"
                            }
                        ))
            else:
                error_msg = page_result.get("error", "Unknown error")
                logger.warning(f"Page {page_num}: OCR failed or no text extracted - {error_msg}")
        
        # Extract and process images from each page with Tesseract OCR
        logger.info(f"Extracting images from {total_pages} pages and applying Tesseract OCR...")
        image_count = 0
        
        for page_num in range(total_pages):
            page = doc[page_num]
            
            # Extract images from this page
            page_images = _extract_images_from_pdf_page(page, page_num + 1, file_name)
            
            # Process each image with Tesseract OCR
            for image_bytes, image_meta in page_images:
                try:
                    image_count += 1
                    logger.info(f"Processing image {image_count} from page {page_num + 1} with Tesseract OCR...")
                    
                    # Apply Tesseract OCR to the image
                    ocr_image_result = extract_text_with_ocr(image_bytes, preprocess=True)
                    
                    if ocr_image_result.get("success") and ocr_image_result.get("text"):
                        image_text = ocr_image_result.get("text", "").strip()
                        image_confidence = ocr_image_result.get("confidence", 0)
                        
                        if image_text:
                            # Create document for OCR'd image text
                            documents.append(Document(
                                page_content=f"Image content (OCR): {image_text}",
                                metadata={
                                    "type": "image_ocr",
                                    "source": file_name,
                                    "file_path": file_path,
                                    "page_number": page_num + 1,
                                    "image_index": image_meta.get("image_index", 0),
                                    "image_format": image_meta.get("format", "unknown"),
                                    "file_type": "pdf",
                                    "ocr_applied": True,
                                    "ocr_engine": "Tesseract",
                                    "ocr_confidence": image_confidence,
                                    "authoritative": True,  # OCR text from images is authoritative
                                    "ingestion_timestamp": datetime.now().isoformat(),
                                    "total_pages": total_pages,
                                    "citation": f"{file_name}, Page {page_num + 1}, Image {image_meta.get('image_index', 0) + 1} (OCR)"
                                }
                            ))
                            logger.info(f"✓ Extracted text from image {image_count}: {len(image_text)} chars")
                        else:
                            logger.warning(f"No text extracted from image {image_count} on page {page_num + 1}")
                    else:
                        # If OCR fails, still create a document noting the image exists
                        error_msg = ocr_image_result.get("error", "Unknown error")
                        logger.warning(f"Tesseract OCR failed for image {image_count} on page {page_num + 1}: {error_msg}")
                        documents.append(Document(
                            page_content=f"Image detected on page {page_num + 1} but OCR extraction failed: {error_msg}",
                            metadata={
                                "type": "image_ocr_failed",
                                "source": file_name,
                                "file_path": file_path,
                                "page_number": page_num + 1,
                                "image_index": image_meta.get("image_index", 0),
                                "file_type": "pdf",
                                "ocr_applied": True,
                                "ocr_engine": "Tesseract",
                                "ocr_error": error_msg,
                                "authoritative": False,
                                "ingestion_timestamp": datetime.now().isoformat(),
                                "total_pages": total_pages,
                                "citation": f"{file_name}, Page {page_num + 1}, Image {image_meta.get('image_index', 0) + 1} (OCR-failed)"
                            }
                        ))
                    
                    # Also generate BLIP description for context (if available)
                    try:
                        image_description = _generate_image_description(image_bytes, image_meta.get("format", "PNG"))
                        if image_description and "unavailable" not in image_description.lower():
                            documents.append(Document(
                                page_content=f"Image description: {image_description}",
                                metadata={
                                    "type": "image_description",
                                    "source": file_name,
                                    "file_path": file_path,
                                    "page_number": page_num + 1,
                                    "image_index": image_meta.get("image_index", 0),
                                    "file_type": "pdf",
                                    "authoritative": False,  # BLIP descriptions are supplementary
                                    "ingestion_timestamp": datetime.now().isoformat(),
                                    "total_pages": total_pages,
                                    "citation": f"{file_name}, Page {page_num + 1}, Image {image_meta.get('image_index', 0) + 1} (BLIP-description)"
                                }
                            ))
                    except Exception as e:
                        logger.debug(f"BLIP description generation skipped for image {image_count}: {str(e)}")
                        
                except Exception as e:
                    logger.warning(f"Failed to process image {image_count} from page {page_num + 1}: {str(e)}")
                    continue
        
        doc.close()
        
        text_chunks = len([d for d in documents if d.metadata.get("type") not in ["image_ocr", "image_ocr_failed", "image_description"]])
        image_ocr_chunks = len([d for d in documents if d.metadata.get("type") == "image_ocr"])
        image_desc_chunks = len([d for d in documents if d.metadata.get("type") == "image_description"])
        
        logger.info(f"✓ Created {len(documents)} OCR document chunks from {file_name}:")
        logger.info(f"  - {text_chunks} text chunks (page OCR, {total_ocr_text} total chars)")
        logger.info(f"  - {image_ocr_chunks} image OCR chunks")
        logger.info(f"  - {image_desc_chunks} image description chunks")
        logger.info(f"  - {image_count} total images processed")
        
        if total_ocr_text < 50:
            logger.warning(f"⚠️  Very little text extracted via OCR ({total_ocr_text} chars)")
            logger.warning(f"   This might indicate: 1) Poor image quality, 2) PDF is mostly images, 3) Incorrect language setting")
        
        if not documents:
            logger.error(f"✗ No documents created from PDF with OCR: {file_name}")
            logger.error(f"   Total OCR text: {total_ocr_text} characters")
            raise RuntimeError(f"OCR processing resulted in no extractable content from {file_name}")
        
        return documents
    
    # Standard native PDF processing (non-OCR mode)
    try:
        if PDF_LIBRARY == "pdfplumber":
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                total_pages = len(pdf.pages)
                for page_num, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text and text.strip():
                        # Chunk the page text
                        chunks = _chunk_text(text, chunk_size, chunk_overlap)
                        for chunk_idx, chunk in enumerate(chunks):
                            documents.append(Document(
                                page_content=chunk,
                                metadata={
                                    "source": file_name,
                                    "file_path": file_path,
                                    "page_number": page_num,
                                    "chunk_index": chunk_idx,
                                    "total_chunks_in_page": len(chunks),
                                    "file_type": "pdf",
                                    "ingestion_timestamp": datetime.now().isoformat(),
                                    "total_pages": total_pages,
                                    "citation": f"{file_name}, Page {page_num}"
                                }
                            ))
        else:
            # Use PyMuPDF (fitz)
            doc = fitz.open(file_path)
            total_pages = len(doc)
            total_text_extracted = 0
            
            for page_num in range(total_pages):
                page = doc[page_num]
                text = page.get_text()
                
                # Log text extraction for debugging
                text_length = len(text.strip()) if text else 0
                logger.debug(f"Page {page_num + 1}: Extracted {text_length} characters")
                
                # Rule 1: If page has text → treat as text (authoritative)
                if text and text.strip():
                    total_text_extracted += text_length
                    # Chunk the page text (unchanged - remains authoritative)
                    chunks = _chunk_text(text, chunk_size, chunk_overlap)
                    for chunk_idx, chunk in enumerate(chunks):
                        documents.append(Document(
                            page_content=chunk,
                            metadata={
                                "source": file_name,
                                "file_path": file_path,
                                "page_number": page_num + 1,
                                "chunk_index": chunk_idx,
                                "total_chunks_in_page": len(chunks),
                                "file_type": "pdf",
                                "ingestion_timestamp": datetime.now().isoformat(),
                                "total_pages": total_pages,
                                "citation": f"{file_name}, Page {page_num + 1}"
                            }
                        ))
                elif text_length <= 10:
                    # Very little text extracted - likely scanned PDF
                    logger.warning(f"Page {page_num + 1}: Only {text_length} characters extracted - PDF may be scanned. Consider using OCR mode.")
                
                # Rule 2 & 3: Extract images from this page and pass to vision model
                page_images = _extract_images_from_pdf_page(page, page_num + 1, file_name)
                
                # Rule 4 & 5: Store each image description as separate Document with metadata
                for image_bytes, image_meta in page_images:
                    try:
                        # Generate image description using vision model (stub for now)
                        image_description = _generate_image_description(
                            image_bytes, 
                            image_meta.get("format", "PNG")
                        )
                        
                        # Create separate Document for image description (non-authoritative)
                        documents.append(Document(
                            page_content=image_description,
                            metadata={
                                "type": "image_description",
                                "authoritative": False,
                                "confidence": "low",
                                "page_number": page_num + 1,
                                "source": file_name,
                                "file_path": file_path,
                                "image_index": image_meta.get("image_index", 0),
                                "image_format": image_meta.get("format", "unknown"),
                                "file_type": "pdf",
                                "ingestion_timestamp": datetime.now().isoformat(),
                                "total_pages": total_pages,
                                "citation": f"{file_name}, Page {page_num + 1}, Image {image_meta.get('image_index', 0) + 1} (non-authoritative)"
                            }
                        ))
                    except Exception as e:
                        logger.warning(f"Failed to generate description for image on page {page_num + 1}: {str(e)}")
                        continue
                            
            doc.close()
            
            # Check if we got very little text - might be a scanned PDF
            # Auto-fallback to OCR if text extraction yields < 50 chars
            if total_text_extracted < 50 and total_pages > 0 and not use_ocr:
                logger.warning(f"⚠️  Very little text extracted ({total_text_extracted} chars from {total_pages} pages)")
                logger.warning(f"   Attempting automatic OCR fallback...")
                
                try:
                    from services.ocr_service import ocr_pdf_document, is_tesseract_installed
                    if is_tesseract_installed():
                        logger.info(f"🔄 Auto-enabling Tesseract OCR for scanned PDF: {file_name}")
                        # Recursively call with OCR enabled
                        return parse_pdf(file_path, chunk_size, chunk_overlap, use_ocr=True)
                    else:
                        logger.warning(f"   Tesseract OCR not available for auto-fallback")
                except Exception as e:
                    logger.warning(f"   OCR fallback failed: {str(e)}")
                    logger.warning(f"   This PDF may be scanned. Try uploading with is_scanned=True to enable OCR")
            
            if not documents:
                logger.error(f"✗ No documents created from PDF: {file_name}")
                logger.error(f"   Total text extracted: {total_text_extracted} characters from {total_pages} pages")
                logger.error(f"   This might be a scanned PDF - try using OCR mode (is_scanned=True)")
        
        logger.info(f"✓ Parsed PDF: {file_name}, {len(documents)} chunks from {total_pages} pages, {total_text_extracted} total chars")
        return documents
        
    except Exception as e:
        logger.error(f"✗ Error parsing PDF {file_path}: {str(e)}")
        raise


def parse_docx(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """
    Parse DOCX file with paragraph-level tracking and smart chunking.
    
    Args:
        file_path: Path to DOCX file
        chunk_size: Number of characters per chunk
        chunk_overlap: Overlap between chunks
    
    Returns:
        List of Document objects with metadata including paragraph numbers for citations
    """
    if not DOCX_AVAILABLE:
        raise ImportError("DOCX parsing not available. Install: pip install python-docx")
    
    documents = []
    file_name = os.path.basename(file_path)
    
    try:
        doc = DocxDocument(file_path)
        
        # Rule 1: Extract text with paragraph tracking (authoritative)
        paragraphs = []
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                paragraphs.append({
                    "text": text,
                    "paragraph_index": para_idx + 1
                })
        
        if not paragraphs:
            logger.warning(f"No text content found in DOCX: {file_path}")
            documents = []
        else:
            # Combine paragraphs and chunk (unchanged - authoritative text)
            full_text = "\n\n".join([p["text"] for p in paragraphs])
            chunks = _chunk_text(full_text, chunk_size, chunk_overlap)
            
            for chunk_idx, chunk in enumerate(chunks):
                # Determine which paragraphs are in this chunk
                start_para = _find_paragraph_for_chunk(chunk, paragraphs, chunk_idx, len(chunks))
                
                documents.append(Document(
                    page_content=chunk,
                    metadata={
                        "source": file_name,
                        "file_path": file_path,
                        "paragraph_start": start_para,
                        "chunk_index": chunk_idx,
                        "total_chunks": len(chunks),
                        "file_type": "docx",
                        "ingestion_timestamp": datetime.now().isoformat(),
                        "total_paragraphs": len(paragraphs),
                        "citation": f"{file_name}, Paragraph ~{start_para}"
                    }
                ))
            
            # Rule 2 & 3: Extract images from DOCX and pass to vision model
            docx_images = _extract_images_from_docx(file_path)
            
            # Rule 4 & 5: Store each image description as separate Document with metadata
            for image_bytes, image_meta in docx_images:
                try:
                    # Generate image description using vision model
                    image_description = _generate_image_description(
                        image_bytes,
                        image_meta.get("format", "PNG")
                    )
                    
                    # Create separate Document for image description (non-authoritative)
                    documents.append(Document(
                        page_content=image_description,
                        metadata={
                            "type": "image_description",
                            "authoritative": False,
                            "confidence": "low",
                            "page_number": 0,  # DOCX doesn't have pages
                            "source": file_name,
                            "file_path": file_path,
                            "image_index": image_meta.get("image_index", 0),
                            "image_format": image_meta.get("format", "unknown"),
                            "image_path": image_meta.get("image_path", ""),
                            "file_type": "docx",
                            "ingestion_timestamp": datetime.now().isoformat(),
                            "citation": f"{file_name}, Image {image_meta.get('image_index', 0) + 1} (non-authoritative)"
                        }
                    ))
                except Exception as e:
                    logger.warning(f"Failed to generate description for image {image_meta.get('image_index', 0)}: {str(e)}")
                    continue
        
        text_chunks = len([d for d in documents if d.metadata.get("type") != "image_description"])
        image_chunks = len([d for d in documents if d.metadata.get("type") == "image_description"])
        logger.info(f"✓ Parsed DOCX: {file_name}, {text_chunks} text chunks from {len(paragraphs) if paragraphs else 0} paragraphs, {image_chunks} image descriptions")
        return documents
        
    except Exception as e:
        logger.error(f"✗ Error parsing DOCX {file_path}: {str(e)}")
        raise


def parse_txt(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """
    Parse TXT file with encoding detection and smart chunking.
    
    Args:
        file_path: Path to TXT file
        chunk_size: Number of characters per chunk
        chunk_overlap: Overlap between chunks
    
    Returns:
        List of Document objects with metadata including line numbers for citations
    """
    documents = []
    file_name = os.path.basename(file_path)
    
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        text = None
        used_encoding = None
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                used_encoding = encoding
                break
            except UnicodeDecodeError:
                continue
        
        if text is None:
            # Fallback to binary read
            with open(file_path, 'rb') as f:
                text = f.read().decode('utf-8', errors='ignore')
            used_encoding = 'utf-8 (with errors ignored)'
        
        if not text or not text.strip():
            logger.warning(f"Empty or unreadable TXT file: {file_path}")
            return []
        
        # Split into lines for better chunking
        lines = text.split('\n')
        chunks = _chunk_text(text, chunk_size, chunk_overlap)
        
        for chunk_idx, chunk in enumerate(chunks):
            # Estimate line range for this chunk
            start_line = _estimate_line_number(chunk, lines, chunk_idx, len(chunks))
            
            documents.append(Document(
                page_content=chunk,
                metadata={
                    "source": file_name,
                    "file_path": file_path,
                    "line_start": start_line,
                    "chunk_index": chunk_idx,
                    "total_chunks": len(chunks),
                    "file_type": "txt",
                    "encoding": used_encoding,
                    "ingestion_timestamp": datetime.now().isoformat(),
                    "total_lines": len(lines),
                    "citation": f"{file_name}, Line ~{start_line}"
                }
            ))
        
        logger.info(f"✓ Parsed TXT: {file_name}, {len(documents)} chunks from {len(lines)} lines (encoding: {used_encoding})")
        return documents
        
    except Exception as e:
        logger.error(f"✗ Error parsing TXT {file_path}: {str(e)}")
        raise


def parse_excel(file_path: str) -> List[Document]:
    """
    Parse Excel file (XLS/XLSX) - enhanced version with better metadata.
    Each row becomes a document with full citation information.
    """
    try:
        df = pd.read_excel(file_path)
        documents = []
        file_name = os.path.basename(file_path)

        for idx, row in df.iterrows():
            content_lines = []

            for col in df.columns:
                value = row[col]
                if pd.notna(value):
                    col_clean = str(col).strip().capitalize()
                    value_clean = str(value).strip()
                    content_lines.append(f"{col_clean}: {value_clean}")

            if content_lines:
                content = "\n".join(content_lines).lower()

                documents.append(Document(
                    page_content=content,
                    metadata={
                        "row_index": idx + 2,  # +2 because Excel is 1-indexed and has header
                        "source": file_name,
                        "file_path": file_path,
                        "columns": list(df.columns),
                        "file_type": "excel",
                        "ingestion_timestamp": datetime.now().isoformat(),
                        "total_rows": len(df),
                        "citation": f"{file_name}, Row {idx + 2}"
                    }
                ))
        
        logger.info(f"✓ Parsed Excel: {file_name}, {len(documents)} rows")
        return documents
        
    except Exception as e:
        logger.error(f"✗ Error parsing Excel {file_path}: {str(e)}")
        raise


def parse_document(file_path: str, chunk_size: int = 1000, chunk_overlap: int = 200, use_ocr: bool = False) -> List[Document]:
    """
    Universal document parser that automatically detects file type and parses accordingly.
    
    Supported formats: PDF, DOCX, DOC, TXT, XLSX, XLS
    Supports OCR for scanned PDFs.
    
    Args:
        file_path: Path to document file
        chunk_size: Number of characters per chunk (for text-based formats)
        chunk_overlap: Overlap between chunks for context preservation
        use_ocr: If True, use OCR for scanned PDFs (default: False)
    
    Returns:
        List of Document objects with comprehensive metadata for citations
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        if not PDF_AVAILABLE:
            raise ImportError("PDF parsing not available. Install: pip install pymupdf")
        return parse_pdf(file_path, chunk_size, chunk_overlap, use_ocr=use_ocr)
    
    elif file_ext in ['.docx', '.doc']:
        if not DOCX_AVAILABLE:
            raise ImportError("DOCX parsing not available. Install: pip install python-docx")
        return parse_docx(file_path, chunk_size, chunk_overlap)
    
    elif file_ext in ['.txt', '.text']:
        return parse_txt(file_path, chunk_size, chunk_overlap)
    
    elif file_ext in ['.xlsx', '.xls']:
        return parse_excel(file_path)
    
    else:
        raise ValueError(f"Unsupported file format: {file_ext}. Supported: .pdf, .docx, .doc, .txt, .xlsx, .xls")


def get_supported_formats() -> List[str]:
    """Return list of supported file formats."""
    formats = ['.xlsx', '.xls', '.txt']
    
    if PDF_AVAILABLE:
        formats.append('.pdf')
    
    if DOCX_AVAILABLE:
        formats.extend(['.docx', '.doc'])
    
    return formats


def check_format_support() -> Dict[str, bool]:
    """Check which formats are currently supported based on installed libraries."""
    return {
        "pdf": PDF_AVAILABLE,
        "docx": DOCX_AVAILABLE,
        "txt": True,
        "excel": True
    }


# Helper functions for multimodal support

# Vision model initialization (lazy loading)
_vision_model = None
_vision_processor = None
VISION_MODEL_AVAILABLE = False

def _initialize_vision_model():
    """
    Initialize BLIP vision model for image captioning (lazy loading).
    Model is loaded only once and cached for reuse.
    """
    global _vision_model, _vision_processor, VISION_MODEL_AVAILABLE
    
    if _vision_model is not None:
        return True
    
    try:
        from transformers import BlipProcessor, BlipForConditionalGeneration
        
        logger.info("Loading BLIP vision model (first time only)...")
        _vision_processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
        _vision_model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
        VISION_MODEL_AVAILABLE = True
        logger.info("✓ BLIP vision model loaded successfully")
        return True
        
    except ImportError:
        logger.warning("transformers not available. Install: pip install transformers")
        VISION_MODEL_AVAILABLE = False
        return False
    except Exception as e:
        logger.warning(f"Failed to load vision model: {str(e)}")
        VISION_MODEL_AVAILABLE = False
        return False


def _generate_image_description(image_data: bytes, image_format: str = "PNG") -> str:
    """
    Generate description of an image using BLIP vision model.
    Falls back to placeholder if model unavailable.
    
    Args:
        image_data: Raw image bytes
        image_format: Image format (PNG, JPEG, etc.)
    
    Returns:
        Text description of the image
    """
    try:
        if not PIL_AVAILABLE:
            return f"[Vision unavailable] Image detected, format: {image_format}"
        
        # Open image
        image = Image.open(io.BytesIO(image_data))
        width, height = image.size
        
        # Try to use BLIP model
        if _initialize_vision_model():
            try:
                # Convert to RGB if needed (BLIP requires RGB)
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Generate caption
                inputs = _vision_processor(image, return_tensors="pt")
                outputs = _vision_model.generate(**inputs, max_length=50)
                description = _vision_processor.decode(outputs[0], skip_special_tokens=True)
                
                logger.debug(f"Generated image description: {description}")
                return f"Image shows: {description} ({width}x{height}px)"
                
            except Exception as e:
                logger.warning(f"Vision model inference failed: {str(e)}")
                return f"Image detected: {width}x{height}px, {image.mode} mode (caption failed)"
        else:
            # Fallback to placeholder
            return f"[Vision model unavailable] Image: {width}x{height}px, {image.mode} mode"
            
    except Exception as e:
        logger.warning(f"Error analyzing image: {str(e)}")
        return f"[Vision error] Image analysis failed: {str(e)}"


def _extract_images_from_pdf_page(page, page_num: int, file_name: str) -> List[Tuple[bytes, Dict[str, Any]]]:
    """
    Extract images from a PDF page using PyMuPDF.
    
    Args:
        page: PyMuPDF page object
        page_num: Page number (1-indexed)
        file_name: Source file name
    
    Returns:
        List of tuples: (image_bytes, image_metadata)
    """
    images = []
    
    try:
        if PDF_LIBRARY == "pymupdf" and hasattr(page, 'get_images'):
            image_list = page.get_images(full=True)
            
            for img_idx, img in enumerate(image_list):
                try:
                    xref = img[0]
                    base_image = page.parent.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    
                    images.append((
                        image_bytes,
                        {
                            "image_index": img_idx,
                            "xref": xref,
                            "format": image_ext,
                            "page_number": page_num,
                            "source": file_name
                        }
                    ))
                except Exception as e:
                    logger.debug(f"Could not extract image {img_idx} from page {page_num}: {str(e)}")
                    continue
        else:
            logger.debug(f"Image extraction not available for {PDF_LIBRARY} on page {page_num}")
            
    except Exception as e:
        logger.debug(f"Error during image extraction from page {page_num}: {str(e)}")
    
    return images


def _extract_images_from_docx(file_path: str) -> List[Tuple[bytes, Dict[str, Any]]]:
    """
    Extract images from a DOCX file.
    DOCX files are ZIP archives - extract from word/media/ folder.
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        List of tuples: (image_bytes, image_metadata)
    """
    images = []
    
    if not DOCX_AVAILABLE:
        return images
    
    try:
        import zipfile
        
        with zipfile.ZipFile(file_path, 'r') as docx_zip:
            file_list = docx_zip.namelist()
            
            # Find images in word/media/ folder
            image_extensions = ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp']
            image_files = [f for f in file_list if f.startswith('word/media/') and 
                          any(f.lower().endswith(ext) for ext in image_extensions)]
            
            for img_idx, img_path in enumerate(image_files):
                try:
                    image_bytes = docx_zip.read(img_path)
                    img_ext = Path(img_path).suffix.lower().replace('.', '')
                    
                    images.append((
                        image_bytes,
                        {
                            "image_index": img_idx,
                            "image_path": img_path,
                            "format": img_ext,
                            "source": os.path.basename(file_path)
                        }
                    ))
                except Exception as e:
                    logger.debug(f"Could not extract image {img_path}: {str(e)}")
                    continue
                    
    except Exception as e:
        logger.debug(f"Error extracting images from DOCX {file_path}: {str(e)}")
    
    return images


def _chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    """
    Split text into chunks with overlap for context preservation.
    Uses sentence boundaries when possible for better semantic coherence.
    """
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        if end >= len(text):
            chunks.append(text[start:].strip())
            break
        
        # Try to break at sentence boundary (., !, ?)
        sentence_end = max(
            text.rfind('.', start, end),
            text.rfind('!', start, end),
            text.rfind('?', start, end)
        )
        
        if sentence_end > start:
            end = sentence_end + 1
        else:
            # Try to break at paragraph boundary
            para_end = text.rfind('\n\n', start, end)
            if para_end > start:
                end = para_end + 2
            else:
                # Try to break at word boundary
                word_end = text.rfind(' ', start, end)
                if word_end > start:
                    end = word_end
        
        chunk_text = text[start:end].strip()
        if chunk_text:
            chunks.append(chunk_text)
        
        start = end - chunk_overlap if chunk_overlap > 0 else end
    
    return chunks


def _find_paragraph_for_chunk(chunk: str, paragraphs: List[Dict], chunk_idx: int, total_chunks: int) -> int:
    """Estimate which paragraph a chunk starts from."""
    if not paragraphs:
        return 1
    
    # Simple estimation based on chunk position
    estimated_para = int((chunk_idx / total_chunks) * len(paragraphs)) + 1
    return min(estimated_para, len(paragraphs))


def _estimate_line_number(chunk: str, lines: List[str], chunk_idx: int, total_chunks: int) -> int:
    """Estimate which line a chunk starts from."""
    if not lines:
        return 1
    
    estimated_line = int((chunk_idx / total_chunks) * len(lines)) + 1
    return min(estimated_line, len(lines))
